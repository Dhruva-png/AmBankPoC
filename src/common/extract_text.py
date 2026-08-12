from __future__ import annotations

import sys
from pathlib import Path


def extract_docx_text(path: str) -> str:
    import docx

    doc = docx.Document(path)
    parts: list[str] = [para.text for para in doc.paragraphs]
    for ti, table in enumerate(doc.tables):
        parts.append(f"\n[TABLE {ti}]")
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            parts.append(" | ".join(cells))
    return "\n".join(parts)


def extract_doc_text(path: str) -> str:
    import win32com.client

    word = win32com.client.DispatchEx("Word.Application")
    word.Visible = False
    try:
        wdoc = word.Documents.Open(str(Path(path).resolve()), ReadOnly=True)
        try:
            return wdoc.Content.Text
        finally:
            wdoc.Close(False)
    finally:
        word.Quit()


def extract_pdf_text(path: str) -> str:
    import pdfplumber

    parts: list[str] = []
    with pdfplumber.open(path) as pdf:
        for i, page in enumerate(pdf.pages):
            text = page.extract_text() or ""
            parts.append(f"--- page {i + 1} ---\n{text}")
    return "\n\n".join(parts)


def pdf_has_text_layer(path: str, min_chars: int = 40) -> bool:
    import pdfplumber

    with pdfplumber.open(path) as pdf:
        total = sum(len((page.extract_text() or "").strip()) for page in pdf.pages)
    return total >= min_chars


def render_pdf_pages_to_images(path: str, out_dir: str, scale: float = 2.0) -> list[str]:
    import pypdfium2 as pdfium

    Path(out_dir).mkdir(parents=True, exist_ok=True)
    stem = Path(path).stem
    pdf = pdfium.PdfDocument(path)
    out_paths = []
    for i in range(len(pdf)):
        page = pdf[i]
        bitmap = page.render(scale=scale)
        image = bitmap.to_pil()
        out_path = str(Path(out_dir) / f"{stem}_p{i + 1:02d}.png")
        image.save(out_path)
        out_paths.append(out_path)
    return out_paths


def render_pdf_first_page(path: str, out_dir: str, scale: float = 1.3) -> str:
    import pypdfium2 as pdfium

    Path(out_dir).mkdir(parents=True, exist_ok=True)
    stem = Path(path).stem
    pdf = pdfium.PdfDocument(path)
    bitmap = pdf[0].render(scale=scale)
    out_path = str(Path(out_dir) / f"{stem}_classify.png")
    bitmap.to_pil().save(out_path)
    return out_path


def extract_text(path: str) -> str:
    suffix = Path(path).suffix.lower()
    if suffix == ".docx":
        return extract_docx_text(path)
    if suffix == ".doc":
        return extract_doc_text(path)
    if suffix == ".pdf":
        return extract_pdf_text(path)
    raise ValueError(f"Unsupported file type: {suffix} ({path})")


def main() -> None:
    if len(sys.argv) < 2:
        print("usage: python extract_text.py <file> [-o out.txt]", file=sys.stderr)
        raise SystemExit(2)
    src = sys.argv[1]
    text = extract_text(src)
    if "-o" in sys.argv:
        out_path = sys.argv[sys.argv.index("-o") + 1]
        Path(out_path).write_text(text, encoding="utf-8")
        print(f"wrote {len(text)} chars to {out_path}")
    else:
        print(text)


if __name__ == "__main__":
    main()
