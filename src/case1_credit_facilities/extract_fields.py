from __future__ import annotations

import re
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parents[1] / "common"))
from extract_text import extract_text  # noqa: E402
from classify import classify_text  # noqa: E402

CATEGORIES = {
    "credit_paper": (
        "Approved Credit Paper -- an internal bank memo containing a 'Principal Terms and "
        "Conditions' table (Customer, Facility Type, Purpose, Security, Covenants & Special "
        "Conditions) and a facility limits table."
    ),
    "letter_of_offer": (
        "Letter of Offer -- a letter addressed to the customer with a 'FACILITY OF RM...' "
        "heading, a 'PURPOSE OF THE FACILITY' clause, and signed for and on behalf of the bank."
    ),
}


def classify_document(path: str) -> str:
    try:
        text = extract_text(path)
    except Exception:
        return "unknown"
    return classify_text(text, CATEGORIES)


def _find_principal_terms_table(doc):
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if len(cells) >= 3 and cells[0] == "Customer" and cells[1] == ":":
                return table
    return None


def _find_basic_info_table(doc):
    for table in doc.tables:
        if len(table.rows) > 1 and "CUSTOMER NAME" in table.rows[1].cells[0].text.upper():
            return table
    return None


def _parse_address_cell(cell_text: str) -> str:
    lines = [ln.strip() for ln in cell_text.split("\n") if ln.strip()]
    if len(lines) < 2:
        return ""
    parts = [lines[1]]
    for ln in lines[2:]:
        if ln.upper().startswith(("POSTCODE", "STATE")):
            parts.append(ln)
    return ", ".join(parts)


def _find_facility_limit_tables(doc):
    tables = []
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if cells[:2] == ["Customer", "Facility"] and "Pricing" in cells:
                tables.append(table)
                break
    return tables


def _header_index(header_cells: list[str], label: str) -> int:
    return header_cells.index(label)


def extract_facility_limits(table) -> list[dict]:
    header = [c.text.strip() for c in table.rows[1].cells]
    idx_facility = _header_index(header, "Facility")
    idx_limit = _header_index(header, "Limit")
    idx_pricing = _header_index(header, "Pricing")
    idx_security = _header_index(header, "Security Details")
    book = table.rows[0].cells[0].text.strip()

    facilities = []
    for row in table.rows[2:]:
        cells = [c.text.strip() for c in row.cells]
        customer = cells[0]
        if customer.lower() in ("total", "grand total", ""):
            continue
        facilities.append(
            {
                "book": book,
                "customer": customer,
                "facility_code": cells[idx_facility],
                "limit_rm000": cells[idx_limit],
                "pricing": cells[idx_pricing],
                "security_details": cells[idx_security],
            }
        )
    return facilities


def extract_credit_paper_fields(docx_path: str) -> dict:
    import docx

    doc = docx.Document(docx_path)
    doc_name = Path(docx_path).name

    terms_table = _find_principal_terms_table(doc)
    if terms_table is None:
        raise ValueError("Could not locate the Principal Terms and Conditions table")

    terms: dict[str, str] = {}
    for row in terms_table.rows:
        cells = [c.text.strip() for c in row.cells]
        if len(cells) >= 3 and cells[1] == ":":
            terms[cells[0]] = cells[2]

    facilities: list[dict] = []
    for limit_table in _find_facility_limit_tables(doc):
        facilities.extend(extract_facility_limits(limit_table))

    registered_address = business_address = ""
    basic_info_table = _find_basic_info_table(doc)
    if basic_info_table is not None and len(basic_info_table.rows) > 2:
        addr_row = basic_info_table.rows[2]
        registered_address = _parse_address_cell(addr_row.cells[0].text)
        business_address = _parse_address_cell(addr_row.cells[1].text)

    terms_source = f"{doc_name} — Principal Terms and Conditions table"
    limits_source = f"{doc_name} — Facility limit table (per book)"
    basic_info_source = f"{doc_name} — Basic Information block"

    return {
        "source_file": str(docx_path),
        "customer": terms.get("Customer", ""),
        "facility_type": terms.get("Facility Type", ""),
        "purpose": terms.get("Purpose", ""),
        "availability_period": terms.get("Availability Period", ""),
        "term_maturity": terms.get("Term/Maturity", ""),
        "fee": terms.get("Fee (Commitment, Facility etc)", ""),
        "security": terms.get("Security (Full details including valuation details)", ""),
        "support_guarantees": terms.get("Support (Guarantees)", ""),
        "conditions_precedent": terms.get("Conditions Precedent", ""),
        "special_conditions": terms.get("Covenants & Special Conditions", ""),
        "registered_address": registered_address,
        "business_address": business_address,
        "facilities": facilities,
        "sources": {
            "customer": terms_source,
            "purpose": terms_source + ", row 'Purpose'",
            "availability_period": terms_source + ", row 'Availability Period'",
            "term_maturity": terms_source + ", row 'Term/Maturity'",
            "security": terms_source + ", row 'Security'",
            "support_guarantees": terms_source + ", row 'Support (Guarantees)'",
            "special_conditions": terms_source + ", row 'Covenants & Special Conditions'",
            "facilities": limits_source,
            "registered_address": basic_info_source + ", 'REGISTERED ADDRESS' row",
            "business_address": basic_info_source + ", 'BUSINESS ADDRESS' row",
        },
    }


_RE_DATE = re.compile(r"Date\s*:?\s*(\d{1,2}\s+\w+\s+\d{4})")
_RE_REF = re.compile(r"Our Ref\.?\s*:?\s*([^\n]+)")
_RE_FACILITY_AMOUNT = re.compile(
    r"FACILITY\s+OF\s+RM\s*([\d,]+)(?:-(\d{2}))?", re.IGNORECASE
)
_RE_ADDRESSEE_BLOCK = re.compile(
    r"(?:\n|^)([A-Z][A-Z0-9 &.,'\-]+(?:SDN\.?\s*BHD\.?|BERHAD))\s*"
    r"\[Registration No\.?\s*([\d]+\s*\([\dA-Z\-]+\))\]",
)
_RE_ADDRESSEE_ADDRESS = re.compile(
    r"\[Registration No\.?\s*[\d]+\s*\([\dA-Z\-]+\)\]\s*\n(.*?)\n\s*\n", re.DOTALL,
)
_RE_ISSUING_ENTITY = re.compile(r"for and on behalf of\s+(AMBANK[A-Z ()]*BERHAD)", re.IGNORECASE)
_RE_PURPOSE = re.compile(
    r"^\s*PURPOSE OF THE FACILITY\s*$\n+(.*?)\n\s*\n",
    re.IGNORECASE | re.MULTILINE | re.DOTALL,
)
_RE_GUARANTOR = re.compile(
    r"Name\s*:?\s*([A-Z][A-Za-z .'\-]+?)\s*MyKad No\.?\s*([\d\-]{12,14})"
)
_RE_SIGNATORY_BLOCK = re.compile(
    r"for and on behalf of\s+AMBANK[A-Z ()]*BERHAD\n+(.*?)ACKNOWLEDGMENT",
    re.IGNORECASE | re.DOTALL,
)


def _extract_signatories(normalized_with_tabs: str) -> list[str]:
    m = _RE_SIGNATORY_BLOCK.search(normalized_with_tabs)
    if not m:
        return []
    lines = [ln.strip() for ln in m.group(1).split("\n") if ln.strip()]
    if not lines:
        return []
    return [n.strip() for n in re.split(r"[ \t]{2,}|\t+", lines[0]) if n.strip()]


def extract_lo_fields(path: str) -> dict:
    text = extract_text(path)
    doc_name = Path(path).name
    normalized = text.replace("\r\n", "\n").replace("\r", "\n")
    signatories = _extract_signatories(normalized)
    flat = re.sub(r"[ \t]+", " ", normalized)

    m = _RE_FACILITY_AMOUNT.search(flat)
    facility_amount = None
    if m:
        cents = m.group(2) or "00"
        facility_amount = f"RM{m.group(1)}.{cents}"

    m = _RE_ADDRESSEE_BLOCK.search(flat)
    addressee_name = m.group(1).strip() if m else ""
    addressee_reg_no = m.group(2).strip() if m else ""

    m = _RE_ADDRESSEE_ADDRESS.search(flat)
    addressee_address = ", ".join(ln.strip() for ln in m.group(1).split("\n") if ln.strip()) if m else ""

    m = _RE_ISSUING_ENTITY.search(flat)
    issuing_entity = m.group(1).strip() if m else ""

    m = _RE_PURPOSE.search(flat)
    purpose = m.group(1).strip() if m else ""

    m = _RE_GUARANTOR.search(flat)
    guarantor_name = m.group(1).strip() if m else ""
    guarantor_nric = m.group(2).strip() if m else ""

    m = _RE_DATE.search(flat)
    letter_date = m.group(1).strip() if m else ""

    m = _RE_REF.search(flat)
    letter_ref = m.group(1).strip() if m else ""

    return {
        "source_file": str(path),
        "raw_text": text,
        "letter_ref": letter_ref,
        "letter_date": letter_date,
        "addressee_name": addressee_name,
        "addressee_registration_no": addressee_reg_no,
        "addressee_address": addressee_address,
        "issuing_entity": issuing_entity,
        "facility_amount": facility_amount,
        "purpose": purpose,
        "guarantor_name": guarantor_name,
        "guarantor_nric": guarantor_nric,
        "signatories": signatories,
        "sources": {
            "facility_amount": f"{doc_name} — 'RE:' facility heading line",
            "addressee_name": f"{doc_name} — addressee block",
            "addressee_address": f"{doc_name} — addressee block",
            "issuing_entity": f"{doc_name} — 'for and on behalf of' signature line",
            "purpose": f"{doc_name} — 'PURPOSE OF THE FACILITY' clause",
            "guarantor_name": f"{doc_name} — guarantor acknowledgment block",
            "letter_date": f"{doc_name} — letter header",
            "signatories": f"{doc_name} — signature block",
        },
    }


def main() -> None:
    import json

    if len(sys.argv) < 2:
        print("usage: python extract_fields.py <credit_paper.docx | lo.doc/.docx/.pdf>", file=sys.stderr)
        raise SystemExit(2)
    path = sys.argv[1]
    if Path(path).suffix.lower() == ".docx":
        try:
            result = extract_credit_paper_fields(path)
        except ValueError:
            result = extract_lo_fields(path)
    else:
        result = extract_lo_fields(path)
    print(json.dumps(result, indent=2, ensure_ascii=False))


if __name__ == "__main__":
    main()
